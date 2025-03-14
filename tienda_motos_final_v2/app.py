from flask import Flask, render_template, request, redirect, url_for, session, flash

import json
import os
from werkzeug.security import generate_password_hash, check_password_hash

from database import create_users_table  # Importa la función que crea la tabla
from reportlab.pdfgen import canvas
from docx import Document
import openpyxl
from datetime import datetime

import sqlite3  # Importa esto para manejar SQLite

DATABASE = 'database.db'

def get_db_connection():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

def create_users_table():
    conn = get_db_connection()
    conn.execute('''
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nombre TEXT NOT NULL,
            email TEXT NOT NULL UNIQUE,
            password TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

# Llamar al crear tabla al iniciar la app
create_users_table()
# ---------------- Configuración Flask ----------------
app = Flask(__name__)
app.secret_key = 'supersecretkey'

MOTOS_FILE = 'motos.json'
NOTIFICACIONES_FILE = 'notificaciones.json'


# Crear tabla de usuarios al iniciar
create_users_table()

# ---------------- Funciones ----------------
def load_motos():
    if os.path.exists(MOTOS_FILE):
        with open(MOTOS_FILE) as file:
            return json.load(file)
    return []


def save_motos(motos):
    with open(MOTOS_FILE, 'w') as file:
        json.dump(motos, file, indent=4)


def load_notificaciones():
    if os.path.exists(NOTIFICACIONES_FILE):
        with open(NOTIFICACIONES_FILE) as file:
            return json.load(file)
    return []


def save_notificaciones(data):
    with open(NOTIFICACIONES_FILE, 'w') as file:
        json.dump(data, file, indent=4)


# ---------------- Rutas Públicas ----------------

@app.route('/')
def index():
    motos = load_motos()
    fecha_actual = datetime.now().strftime('%Y-%m-%d')
    ofertas = [m for m in motos if m['descuento'] > 0 and m['descuento_hasta'] and m['descuento_hasta'] >= fecha_actual]
    return render_template('index.html', ofertas=ofertas, fecha_actual=fecha_actual)


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if request.form['username'] == 'admin' and request.form['password'] == 'admin123':
            session['admin'] = True
            return redirect(url_for('admin'))
        return render_template('login.html', error="Credenciales incorrectas.")
    return render_template('login.html')


# LOGOUT ADMIN
@app.route('/logout')
def logout():
    session.pop('admin', None)  # Elimina solo la sesión del admin
    flash('Sesión de administrador cerrada.', 'info')
    return redirect(url_for('index'))

    # Login de usuario normal
@app.route('/login_usuario', methods=['GET', 'POST'])
def login_usuario():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']

        conn = get_db_connection()
        user = conn.execute('SELECT * FROM usuarios WHERE email = ?', (email,)).fetchone()
        conn.close()

        if user and check_password_hash(user['password'], password):
            session['user_id'] = user['id']
            session['user_name'] = user['nombre']
            flash(f'Bienvenido, {user['nombre']}', 'success')
            return redirect(url_for('index'))  # Redirige al inicio
        else:
            flash('Correo o contraseña incorrectos.', 'danger')

    # 🔴 Este return debe ir DENTRO de la función, NO fuera.
    return render_template('login_usuario.html')



# ----------------------------------------------------------------
# LOGOUT USUARIO
@app.route('/logout_usuario')
def logout_usuario():
    session.pop('user_id', None)  # Elimina solo la sesión del usuario
    session.pop('user_name', None)
    flash('Has cerrado sesión correctamente.', 'info')
    return redirect(url_for('index'))


@app.route('/product')
@app.route('/product')
def catalog():
    motos = load_motos()  # Cargar todas las motos
    query = request.args.get('q', '').strip().lower()  # Obtener parámetro de búsqueda
    fecha_actual = datetime.now().strftime('%Y-%m-%d')  # Fecha actual para validación de descuentos

    if query:  # Si se escribió algo
        motos = [m for m in motos if query in m['nombre'].lower() or query in m['descripcion'].lower()]  # Filtrar

    return render_template('product.html', motos=motos, fecha_actual=fecha_actual)


@app.route('/product/<int:moto_id>')
def product_detail(moto_id):
    motos = load_motos()
    moto = next((m for m in motos if m['id'] == moto_id), None)
    fecha_actual = datetime.now().strftime('%Y-%m-%d')  # 👈 Definir fecha actual para comparaciones
    return render_template('product_detail.html', moto=moto, fecha_actual=fecha_actual) if moto else ("Moto no encontrada", 404)


@app.route('/admin/historial/delete/<int:moto_id>', methods=['POST'])
def delete_historial_moto(moto_id):
    if not session.get('admin'):
        return redirect(url_for('login'))
    
    notificaciones = load_notificaciones()
    # Filtra todas las notificaciones excepto las de esta moto
    nuevas_notificaciones = [n for n in notificaciones if n.get('moto_id') != moto_id]
    
    save_notificaciones(nuevas_notificaciones)
    return redirect(url_for('historial_moto', moto_id=moto_id))

@app.route('/toggle_favorite/<int:moto_id>')
def toggle_favorite(moto_id):
    favoritos = session.get('favoritos', [])
    if moto_id in favoritos:
        favoritos.remove(moto_id)
    else:
        favoritos.append(moto_id)
    session['favoritos'] = favoritos
    return redirect(request.referrer or url_for('catalog'))


@app.route('/favorites')
def favorites():
    favoritos = session.get('favoritos', [])
    motos = [m for m in load_motos() if m['id'] in favoritos]
    return render_template('favorites.html', motos=motos)


@app.route('/compare', methods=['POST'])
def compare():
    ids = request.form.getlist('compare')
    motos = [m for m in load_motos() if str(m['id']) in ids]
    return render_template('compare.html', motos=motos)

  
# ---------------- REPORTES PDF, WORD, EXCEL ----------------
@app.route('/reporte_pdf')
def reporte_pdf():
    motos = load_motos()
    c = canvas.Canvas("static/reporte_motos.pdf")
    y = 800
    c.drawString(100, y, "Reporte de Motos")
    y -= 40
    for m in motos:
        c.drawString(100, y, f"{m['nombre']} - ${m['precio']}")
        y -= 20
    c.save()
    return send_file("static/reporte_motos.pdf", as_attachment=True)


@app.route('/reporte_word')
def reporte_word():
    motos = load_motos()
    doc = Document()
    doc.add_heading('Reporte de Motos', 0)
    for m in motos:
        doc.add_paragraph(f"{m['nombre']} - ${m['precio']}")
    doc.save('static/reporte_motos.docx')
    return send_file('static/reporte_motos.docx', as_attachment=True)


@app.route('/reporte_excel')
def reporte_excel():
    motos = load_motos()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['ID', 'Nombre', 'Precio', 'Descripción'])
    for m in motos:
        ws.append([m['id'], m['nombre'], m['precio'], m['descripcion']])
    wb.save('static/reporte_motos.xlsx')
    return send_file('static/reporte_motos.xlsx', as_attachment=True)

    # Registro de usuario
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        nombre = request.form['nombre']
        email = request.form['email']
        password = request.form['password']
        hashed_password = generate_password_hash(password)  # Aquí ya funciona el hash

        conn = get_db_connection()
        try:
            conn.execute('INSERT INTO usuarios (nombre, email, password) VALUES (?, ?, ?)', 
                         (nombre, email, hashed_password))
            conn.commit()
            flash('¡Usuario registrado exitosamente! Inicia sesión.', 'success')
            return redirect(url_for('login_usuario'))
        except sqlite3.IntegrityError:
            flash('Este correo ya está registrado. Prueba con otro.', 'danger')
        finally:
            conn.close()

    return render_template('register.html')



    


# ---------------- HISTORIAL AGENDAMIENTOS ----------------
@app.route('/admin/historial/<int:moto_id>')
def historial_moto(moto_id):
    notificaciones = load_notificaciones()
    # ✅ Corrección: verificar que 'moto_id' exista
    historial = [n for n in notificaciones if n.get('moto_id') == moto_id]
    return render_template('historial.html', historial=historial)

# ---------------- Agendar Citas ----------------


@app.route('/schedule/<int:moto_id>', methods=['GET', 'POST'])
def schedule_appointment(moto_id):
    motos = load_motos()
    moto = next((m for m in motos if m['id'] == moto_id), None)
    if not moto:
        return "Moto no encontrada", 404

    if request.method == 'POST':
        mensaje = {
            "moto_id": moto_id,
            "moto": moto['nombre'],
            "nombre": request.form['nombre'],
            "email": request.form['email'],
            "telefono": request.form['telefono'],
            "fecha": request.form['fecha'],
            "runt": request.form['runt'],
            "estado": "nuevo"
        }
        notificaciones = load_notificaciones()
        notificaciones.append(mensaje)
        save_notificaciones(notificaciones)
        return redirect(url_for('catalog'))  # ✅ Redirigir al catálogo (product)

    return render_template('schedule.html', moto=moto)

# ---------------- Panel Administrador ----------------
@app.route('/admin', methods=['GET', 'POST'])
def admin():
    if not session.get('admin'):
        return redirect(url_for('login'))

    motos = load_motos()
    notificaciones = load_notificaciones()
    filtro = request.args.get('filtro', '')
    page = int(request.args.get('page', 1))
    per_page = 6

    # Conteo notificaciones
    count_notificaciones = sum(1 for n in notificaciones if n['estado'] == 'nuevo')

    # Conteo agendadas
    for moto in motos:
        moto["agendada_veces"] = sum(1 for n in notificaciones if n.get("moto_id") == moto["id"])

    # Filtro
    if filtro == 'agendada':
        motos = [m for m in motos if m["agendada_veces"] > 0]
    elif filtro == 'disponible':
        motos = [m for m in motos if m["agendada_veces"] == 0]

    # Paginación
    total_pages = (len(motos) + per_page - 1) // per_page
    motos_pag = motos[(page - 1) * per_page: page * per_page]

    # Agregar motos con descuento
    if request.method == 'POST':
        imagen_file = request.files['imagen']
        imagen_filename = imagen_file.filename
        imagen_file.save(os.path.join('static/images', imagen_filename))

        nueva_moto = {
            "id": max([m['id'] for m in motos], default=0) + 1,
            "nombre": request.form['nombre'],
            "precio": float(request.form['precio']),
            "imagen": imagen_filename,
            "descripcion": request.form['descripcion'],
            "descuento": float(request.form.get('descuento', 0)),
            "descuento_hasta": request.form.get('descuento_hasta', '')
        }
        motos.append(nueva_moto)
        save_motos(motos)
        return redirect(url_for('admin'))

    return render_template('admin.html',
                           motos=motos_pag,
                           count_notificaciones=count_notificaciones,
                           total_pages=total_pages,
                           current_page=page,
                           filtro=filtro,
                           now=datetime.now())


@app.route('/admin/remove_discount/<int:moto_id>', methods=['POST'])
def remove_discount(moto_id):
    motos = load_motos()
    for moto in motos:
        if moto['id'] == moto_id:
            moto['descuento'] = 0
            moto['descuento_hasta'] = ""
    save_motos(motos)
    return redirect(url_for('admin'))


@app.route('/admin/edit/<int:moto_id>', methods=['GET', 'POST'])
def edit_moto(moto_id):
    if not session.get('admin'):
        return redirect(url_for('login'))
    motos = load_motos()
    moto = next((m for m in motos if m['id'] == moto_id), None)
    if request.method == 'POST':
        moto["nombre"] = request.form['nombre']
        moto["precio"] = float(request.form['precio'])
        moto["descripcion"] = request.form['descripcion']

        imagen_file = request.files.get('imagen')
        if imagen_file and imagen_file.filename != '':
            imagen_filename = imagen_file.filename
            imagen_file.save(os.path.join('static/images', imagen_filename))
            moto["imagen"] = imagen_filename

        save_motos(motos)
        return redirect(url_for('admin'))
    return render_template('edit_moto.html', moto=moto)


@app.route('/admin/delete/<int:moto_id>', methods=['POST'])
def delete_moto(moto_id):
    if not session.get('admin'):
        return redirect(url_for('login'))
    motos = load_motos()
    motos = [m for m in motos if m['id'] != moto_id]
    save_motos(motos)
    return redirect(url_for('admin'))


# ---------------- Notificaciones ----------------
@app.route('/check_notificaciones')
def check_notificaciones():
    notificaciones = load_notificaciones()
    nuevas = sum(1 for n in notificaciones if n['estado'] == 'nuevo')
    return jsonify({"count": nuevas})


@app.route('/notificaciones')
def notificaciones():
    if not session.get('admin'):
        return redirect(url_for('login'))
    mensajes = load_notificaciones()
    return render_template('notificaciones.html', mensajes=mensajes)


@app.route('/notificaciones/marcar_leido')
def marcar_notificaciones():
    mensajes = load_notificaciones()
    for msg in mensajes:
        msg['estado'] = 'leido'
    save_notificaciones(mensajes)
    return redirect(url_for('notificaciones'))


    # ------------------ Carrito de Compras para Usuarios Registrados ------------------ #

# Añadir moto al carrito (Agendar Cita)
@app.route('/agendar_cita/<int:moto_id>')
def agendar_cita(moto_id):
    if 'user_id' not in session:
        flash('Debes iniciar sesión para agendar una cita.', 'warning')
        return redirect(url_for('login_usuario'))

    # Cargar motos
    motos = load_motos()
    moto = next((m for m in motos if m['id'] == moto_id), None)

    if not moto:
        flash('Moto no encontrada.', 'danger')
        return redirect(url_for('catalog'))

    # Inicializar carrito si no existe
    if 'carrito' not in session:
        session['carrito'] = []

    # Verificar si ya está en el carrito
    if moto_id in session['carrito']:
        flash('Esta moto ya está en tu carrito.', 'info')
    else:
        session['carrito'].append(moto_id)
        flash(f'Moto "{moto["nombre"]}" añadida al carrito para agendar.', 'success')

    return redirect(url_for('catalog'))

# Ver carrito
@app.route('/carrito')
def ver_carrito():
    if 'user_id' not in session:
        flash('Debes iniciar sesión para ver tu carrito.', 'warning')
        return redirect(url_for('login_usuario'))

    carrito_ids = session.get('carrito', [])
    motos = load_motos()
    motos_en_carrito = [m for m in motos if m['id'] in carrito_ids]

    return render_template('carrito.html', motos=motos_en_carrito)

# Eliminar moto del carrito
@app.route('/eliminar_carrito/<int:moto_id>')
def eliminar_carrito(moto_id):
    if 'user_id' not in session:
        flash('Debes iniciar sesión.', 'warning')
        return redirect(url_for('login_usuario'))

    if 'carrito' in session and moto_id in session['carrito']:
        session['carrito'].remove(moto_id)
        flash('Moto eliminada del carrito.', 'info')
    return redirect(url_for('ver_carrito'))

# Vaciar carrito
@app.route('/vaciar_carrito')
def vaciar_carrito():
    if 'user_id' not in session:
        flash('Debes iniciar sesión.', 'warning')
        return redirect(url_for('login_usuario'))

    session['carrito'] = []
    flash('Carrito vaciado.', 'info')
    return redirect(url_for('ver_carrito'))


# FORMULARIO PARA AGENDAR TODAS LAS MOTOS DEL CARRITO
@app.route('/agendar_carrito', methods=['GET', 'POST'])
def agendar_carrito():
    if 'user_id' not in session:
        flash('Debes iniciar sesión para agendar.', 'warning')
        return redirect(url_for('login_usuario'))

    carrito_ids = session.get('carrito', [])
    motos = [m for m in load_motos() if m['id'] in carrito_ids]

    if request.method == 'POST':
        nombre = request.form['nombre']
        email = request.form['email']
        telefono = request.form['telefono']
        tiene_runt = request.form['tiene_runt']
        mensaje = request.form['mensaje']

        notificaciones = load_notificaciones()

        for moto in motos:
            nueva_notificacion = {
                'id': len(notificaciones) + 1,
                'nombre': nombre,
                'email': email,
                'telefono': telefono,
                'mensaje': mensaje,
                'moto_id': moto['id'],
                'moto_nombre': moto['nombre'],
                'estado': 'nuevo',
                'tiene_runt': tiene_runt
            }
            notificaciones.append(nueva_notificacion)

        save_notificaciones(notificaciones)
        session['carrito'] = []  # Limpiar carrito tras agendar
        flash('¡Cita agendada exitosamente para las motos seleccionadas!', 'success')
        return redirect(url_for('index'))

    return render_template('agendar_carrito.html', motos=motos)


# ---------------- Dashboard ----------------
@app.route('/dashboard')
def dashboard():
    if not session.get('admin'):
        return redirect(url_for('login'))
    motos = load_motos()
    notificaciones = load_notificaciones()
    conteo = {m['nombre']: sum(1 for n in notificaciones if n.get('moto_id') == m['id']) for m in motos}
    return render_template('dashboard.html', conteo=conteo)



# ---------------- Contacto ----------------
@app.route('/contacto', methods=['GET', 'POST'])
def contacto():
    if request.method == 'POST':
        # Código de contacto
        pass
    return render_template('contacto.html')


    


# ---------------- Iniciar App ----------------
if __name__ == '__main__':
    app.run(debug=True)
